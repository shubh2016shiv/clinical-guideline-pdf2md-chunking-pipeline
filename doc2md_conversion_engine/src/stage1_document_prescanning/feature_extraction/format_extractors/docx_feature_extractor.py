"""
stage1_document_prescanning/feature_extraction/format_extractors/docx_feature_extractor.py
==========================================================================================
Stage 1 · Step 2 of 3 — the reader for Word (.docx) files.

This module does not try to understand the clinical meaning of a Word document.
It only collects factual evidence that is cheap to inspect:

1. Is there native text?
2. Are there tables?
3. Are there embedded images?
4. Are there figure captions that confirm a visual was intentionally labelled?

Because DOCX is a ZIP archive whose XML carries no rendering coordinates, this
module cannot report page numbers, bounding boxes, or image sizes.  Page count
is *estimated* from character count — see the module note below for why.

Thresholds come from ``settings.yaml`` via ``DocumentFeatureExtractionConfig``.

-----

What is a DOCX file?
--------------------
Despite looking like a single file, a .docx is actually a ZIP archive
containing a folder of XML files.  The body text lives in word/document.xml.
Images are stored as separate files inside the ZIP (e.g. word/media/image1.png)
and connected to the XML via a "relationships table" — essentially a manifest
that says "this image tag in the XML refers to that file in the ZIP".

Word exposes images through two separate APIs because of this split structure.
We need both APIs to get an accurate image count — see
``_count_images_anchored_in_text_flow`` and ``_count_image_files_in_zip_archive``.

Why are page numbers estimated rather than read from the file?
  A DOCX file stores no page information.  Pages are calculated by the Word
  rendering engine at display time and depend on font, margins, zoom level,
  and screen size.  We estimate from character count as a practical approximation.
"""

from __future__ import annotations

import logging
import math
from dataclasses import dataclass, field
from pathlib import Path

from ....contracts.configurations.pipeline_config import (
    DocumentFeatureExtractionConfig,
    DocxFeatureExtractionConfig,
    EngineNeedsEvaluatorConfig,
)
from ....contracts.exceptions import DocumentError
from ...engine_routing.engine_format_compatibility import get_engine_format_compatibility
from ...engine_routing.document_requirements_resolver import resolve_document_requirements
from ..feature_evidence_models import (
    DocumentFeatureProfile,
    FeatureDocumentType,
    LayoutEvidence,
    TableEvidence,
    TextEvidence,
    VisualCandidate,
    VisualCandidateKind,
    VisualEvidence,
)
from ..visual_caption_detector import compact_text, contains_figure_caption, count_figure_caption_lines

logger = logging.getLogger(__name__)


@dataclass
class DocxFeatureTotals:
    """Running totals collected while reading a Word document."""

    total_characters: int = 0
    body_paragraphs: list[str] = field(default_factory=list)

    number_of_tables: int = 0
    max_table_column_count: int = 0
    has_merged_cells: bool = False
    has_nested_tables: bool = False

    column_count: int = 1
    has_floating_text_boxes: bool = False

    images_in_text_flow: int = 0
    images_in_zip_archive: int = 0

    number_of_captioned_figures: int = 0


# ---------------------------------------------------------------------------
# Public entry-point
# ---------------------------------------------------------------------------


def extract_docx_features(
    path: Path,
    config: DocumentFeatureExtractionConfig | None = None,
) -> DocumentFeatureProfile:
    """
    Extract deterministic feature evidence from one Word document.

    The flow is intentionally linear:
    open the document, collect text, count structural elements, then build the
    final ``DocumentFeatureProfile`` consumed by the capability router.

    Step 1: Open the file.
    Step 2: Read all text (body paragraphs + text inside table cells).
    Step 3: Estimate page count from character count.
    Step 4: Count tables.
    Step 5: Count images via both Word APIs (text-flow and ZIP archive).
    Step 6: Count figure captions.
    Step 7: Assemble and return the feature profile.
    """
    resolved_config = config or DocumentFeatureExtractionConfig()
    docx_settings = resolved_config.docx
    evaluator_settings = resolved_config.engine_needs_evaluator
    doc = open_word_document(path)
    totals = DocxFeatureTotals()

    collect_text_from_document(doc, totals)

    estimated_pages = estimate_page_count(totals.total_characters, docx_settings)

    collect_table_evidence_from_document(doc, totals)
    collect_layout_evidence_from_document(doc, totals)
    collect_image_evidence_from_document(doc, totals)
    collect_caption_evidence(totals)

    logger.debug(
        "DOCX feature extraction complete: path=%s estimated_pages=%d "
        "chars=%d tables=%d images=%d captions=%d",
        path.name,
        estimated_pages,
        totals.total_characters,
        totals.number_of_tables,
        max(totals.images_in_text_flow, totals.images_in_zip_archive),
        totals.number_of_captioned_figures,
    )

    return build_docx_feature_profile(
        estimated_pages=estimated_pages,
        totals=totals,
        settings=docx_settings,
        evaluator_settings=evaluator_settings,
    )


# ---------------------------------------------------------------------------
# Step implementations — called in entry-point order
# ---------------------------------------------------------------------------


def collect_text_from_document(doc: object, totals: DocxFeatureTotals) -> None:
    """
    Step 2: Read all text — body paragraphs and text inside table cells.

    Iterating only ``doc.paragraphs`` silently misses all text that sits inside
    a table cell, so we collect both explicitly.
    """
    import docx as _docx  # noqa: PLC0415 — deferred; python-docx may not be installed

    real_doc: _docx.Document = doc  # type: ignore[assignment]

    body_paragraphs: list[str] = [
        paragraph.text
        for paragraph in real_doc.paragraphs
        if paragraph.text.strip()
    ]
    table_cell_paragraphs: list[str] = [
        paragraph.text
        for table in real_doc.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        if paragraph.text.strip()
    ]

    totals.body_paragraphs = body_paragraphs
    all_text = "\n".join([*body_paragraphs, *table_cell_paragraphs])
    totals.total_characters = len(all_text)


def estimate_page_count(total_characters: int, settings: DocxFeatureExtractionConfig) -> int:
    """
    Step 3: Estimate page count from character count.

    ``settings.average_characters_per_page`` comes from
    ``document_feature_extraction.docx.average_characters_per_page`` in
    ``settings.yaml``.  We err toward over-counting pages — safer for downstream
    rendering decisions.
    """
    return max(1, math.ceil(total_characters / settings.average_characters_per_page))


def collect_table_evidence_from_document(doc: object, totals: DocxFeatureTotals) -> None:
    """
    Step 4: Count Word table objects and read their structural complexity.

    DOCX declares table structure directly in XML, so merged cells, nested
    tables, and column counts are exact reads — no rendering or geometry needed.
    These are the signals that decide whether a grid-naive parser will mangle
    the table (→ MinerU) or handle it cleanly (→ Docling).
    """
    from docx.oxml.ns import qn  # noqa: PLC0415 — deferred; python-docx may not be installed

    import docx as _docx  # noqa: PLC0415

    real_doc: _docx.Document = doc  # type: ignore[assignment]
    tables = real_doc.tables
    totals.number_of_tables = len(tables)

    for table in tables:
        totals.max_table_column_count = max(
            totals.max_table_column_count, len(table.columns)
        )
        if not totals.has_nested_tables and _table_has_nested_table(table):
            totals.has_nested_tables = True
        if not totals.has_merged_cells and _table_has_merged_cells(table, qn):
            totals.has_merged_cells = True


def _table_has_nested_table(table: object) -> bool:
    """Return True when any cell of *table* contains another table."""
    return any(cell.tables for row in table.rows for cell in row.cells)


def _table_has_merged_cells(table: object, qn: object) -> bool:
    """
    Return True when any cell carries horizontal (gridSpan) or vertical (vMerge)
    merge geometry — both are single XML attribute checks in the cell's <w:tc>.
    """
    horizontal_merge_tag = qn("w:gridSpan")  # type: ignore[operator]
    vertical_merge_tag = qn("w:vMerge")  # type: ignore[operator]
    for row in table.rows:
        for cell in row.cells:
            cell_xml = cell._tc
            if cell_xml.find(f".//{horizontal_merge_tag}") is not None:
                return True
            if cell_xml.find(f".//{vertical_merge_tag}") is not None:
                return True
    return False


def collect_layout_evidence_from_document(doc: object, totals: DocxFeatureTotals) -> None:
    """
    Read multi-column sections and floating text boxes from the document XML.

    ``<w:cols w:num="N">`` inside a section declares N text columns — a direct,
    exact read with no geometry.  ``<w:txbxContent>`` marks a floating text box,
    which breaks linear reading order.  Both promote to a layout-aware engine.
    """
    from docx.oxml.ns import qn  # noqa: PLC0415

    import docx as _docx  # noqa: PLC0415

    real_doc: _docx.Document = doc  # type: ignore[assignment]

    column_number_attr = qn("w:num")
    max_columns = 1
    for section in real_doc.sections:
        cols_element = section._sectPr.find(qn("w:cols"))
        if cols_element is None:
            continue
        declared = cols_element.get(column_number_attr)
        if declared and declared.isdigit():
            max_columns = max(max_columns, int(declared))
    totals.column_count = max_columns

    totals.has_floating_text_boxes = (
        next(real_doc.element.iter(qn("w:txbxContent")), None) is not None
    )


def collect_image_evidence_from_document(doc: object, totals: DocxFeatureTotals) -> None:
    """
    Step 5: Count images via both Word APIs.

    ``images_in_text_flow`` counts inline shapes; ``images_in_zip_archive``
    counts every image file in the DOCX ZIP.  We record both so the caller can
    log the breakdown and take the higher of the two.
    """
    totals.images_in_text_flow = _count_images_anchored_in_text_flow(doc)
    totals.images_in_zip_archive = _count_image_files_in_zip_archive(doc)


def collect_caption_evidence(totals: DocxFeatureTotals) -> None:
    """
    Step 6: Count figure captions across all body paragraphs.

    A caption (e.g. "Figure 3: Sales by region") confirms that a nearby image
    or chart was intentionally labelled by the author.
    """
    all_body_text = "\n".join(totals.body_paragraphs)
    totals.number_of_captioned_figures = count_figure_caption_lines(all_body_text)


# ---------------------------------------------------------------------------
# Profile assembly
# ---------------------------------------------------------------------------


def build_docx_feature_profile(
    *,
    estimated_pages: int,
    totals: DocxFeatureTotals,
    settings: DocxFeatureExtractionConfig,
    evaluator_settings: EngineNeedsEvaluatorConfig,
) -> DocumentFeatureProfile:
    """
    Translate raw extraction totals into the structured profile consumed by the
    capability router downstream.

    ``DocumentFeatureProfile`` is the single object that leaves this module.
    Everything else here is an intermediate view of the same raw numbers,
    shaped into the three distinct contracts the router expects:

    - *evidence* objects (TextEvidence, TableEvidence, VisualEvidence) — aggregate
      statistics (counts, densities, flags).  The router reads these to decide
      *whether* a capability is needed.

    - *visual_candidates* — a short ranked list of specific elements (an image
      block, a table) worth showing to a vision model for a second opinion.
      These carry enough context (nearby caption, source API) for a model to make
      a meaningful judgment.  Capped by ``settings.max_visual_candidates`` so the
      downstream payload stays bounded.

    - *requirements* — inferred capability flags (e.g. needs_ocr, needs_vlm)
      derived from the evidence above.  The router uses these to filter the
      engine candidates it will score.
    """
    # DOCX image count is the higher of two independent API counts (inline
    # shapes vs. ZIP-archive image parts) — see collect_image_evidence_from_document.
    number_of_images = max(totals.images_in_text_flow, totals.images_in_zip_archive)

    # Aggregate statistics — one object per feature dimension.
    text_evidence = build_text_evidence(totals, estimated_pages)
    table_evidence = build_table_evidence(totals, estimated_pages)
    layout_evidence = build_layout_evidence(totals)
    visual_evidence = build_visual_evidence(totals, estimated_pages, number_of_images)

    # Representative examples — one candidate per feature type (images, tables),
    # not one per instance.  DOCX has no coordinates, so candidates use a
    # document-level label instead of a page number or bounding box.
    visual_candidates = build_visual_candidates(
        body_paragraphs=totals.body_paragraphs,
        number_of_images=number_of_images,
        images_found_in_text_flow=totals.images_in_text_flow,
        images_found_in_zip_archive=totals.images_in_zip_archive,
        number_of_tables=totals.number_of_tables,
        max_candidates=settings.max_visual_candidates,
    )

    # Capability flags — derived from the evidence above, not from the raw totals.
    # resolve_document_requirements encodes the routing heuristics so this function stays
    # focused on assembly.
    requirements = resolve_document_requirements(
        text=text_evidence,
        tables=table_evidence,
        layout=layout_evidence,
        visuals=visual_evidence,
        visual_candidates=visual_candidates,
        settings=evaluator_settings,
    )

    return DocumentFeatureProfile(
        file_type=FeatureDocumentType.DOCX,
        page_or_unit_count=estimated_pages,
        text=text_evidence,
        tables=table_evidence,
        layout=layout_evidence,
        visuals=visual_evidence,
        visual_candidates=visual_candidates,
        format_support=get_engine_format_compatibility(FeatureDocumentType.DOCX),
        requirements=requirements,
    )


def build_text_evidence(totals: DocxFeatureTotals, estimated_pages: int) -> TextEvidence:
    """Summarize native text evidence across the Word document."""
    return TextEvidence(
        total_characters=totals.total_characters,
        pages_or_units_with_text=estimated_pages if totals.total_characters else 0,
        estimated_text_density=totals.total_characters / max(estimated_pages, 1),
        native_text_available=totals.total_characters > 0,
    )


def build_table_evidence(totals: DocxFeatureTotals, estimated_pages: int) -> TableEvidence:
    """Summarize table evidence across the Word document."""
    return TableEvidence(
        count=totals.number_of_tables,
        # Assume at most one table per estimated page — a safe upper bound.
        pages_or_units_with_tables=min(totals.number_of_tables, estimated_pages),
        # DOCX stores no size information for tables in its XML.  We cannot tell
        # a large data table from a small inline one without rendering.
        large_count=0,
        # Structure, by contrast, is declared exactly in the XML.
        max_column_count=totals.max_table_column_count,
        has_merged_cells=totals.has_merged_cells,
        has_nested_tables=totals.has_nested_tables,
    )


def build_layout_evidence(totals: DocxFeatureTotals) -> LayoutEvidence:
    """Summarize page-layout evidence read from the document section XML."""
    return LayoutEvidence(
        column_count=totals.column_count,
        has_floating_text_boxes=totals.has_floating_text_boxes,
    )


def build_visual_evidence(
    totals: DocxFeatureTotals,
    estimated_pages: int,
    number_of_images: int,
) -> VisualEvidence:
    """Summarize image and caption evidence across the Word document."""
    return VisualEvidence(
        embedded_image_count=number_of_images,
        # DOCX XML carries no bounding-box data for images — size only known after rendering.
        large_embedded_image_count=0,
        # Word vector shapes (EMF/WMF) are not counted here.
        vector_graphics_count=0,
        # Chart objects need DrawingML parsing; out of scope here.
        chart_count=0,
        svg_count=0,
        pages_or_units_with_visuals=min(number_of_images, estimated_pages),
        captioned_visual_count=totals.number_of_captioned_figures,
    )


def build_visual_candidates(
    *,
    body_paragraphs: list[str],
    number_of_images: int,
    images_found_in_text_flow: int,
    images_found_in_zip_archive: int,
    number_of_tables: int,
    max_candidates: int,
) -> list[VisualCandidate]:
    """
    Build a short list of visual candidates for optional downstream review.

    ``max_candidates`` comes from
    ``document_feature_extraction.docx.max_visual_candidates`` in
    ``settings.yaml``.

    DOCX has no spatial layout at extraction time — no page coordinates and no
    bounding boxes — so candidates carry a document-level location label instead
    of a page number or area ratio.  One candidate is created per feature type
    (images, tables) rather than per individual instance; the evidence string
    carries the count.
    """
    candidates: list[VisualCandidate] = []

    first_figure_caption = next(
        (compact_text(text) for text in body_paragraphs if contains_figure_caption(text)),
        None,
    )

    if number_of_images:
        candidates.append(
            VisualCandidate(
                kind=VisualCandidateKind.EMBEDDED_IMAGE,
                location_label="docx embedded media",
                caption_or_alt_text=first_figure_caption,
                nearby_text=first_figure_caption,
                evidence=[
                    f"{number_of_images} image(s) found "
                    f"(text-flow count: {images_found_in_text_flow}, "
                    f"zip-archive count: {images_found_in_zip_archive})"
                ],
            )
        )

    if number_of_tables:
        candidates.append(
            VisualCandidate(
                kind=VisualCandidateKind.TABLE,
                location_label="docx tables",
                evidence=[f"{number_of_tables} Word table object(s) found"],
            )
        )

    return candidates[:max_candidates]


# ---------------------------------------------------------------------------
# Image counting helpers
# ---------------------------------------------------------------------------


def _count_images_anchored_in_text_flow(doc: object) -> int:
    """Count images inserted inline into the body text (the most common case)."""
    try:
        return len(doc.inline_shapes)  # type: ignore[union-attr]
    except Exception:
        logger.debug("doc.inline_shapes unavailable; skipping text-flow image count")
        return 0


def _count_image_files_in_zip_archive(doc: object) -> int:
    """
    Count distinct image files stored inside the DOCX ZIP archive.

    ``doc.part.related_parts`` is the relationships table — a dictionary
    mapping each embedded file's internal ID to the file itself.
    We filter it for entries whose content_type (MIME type) starts with
    "image/" to count all image assets, wherever they appear in the document.
    """
    try:
        all_embedded_files = doc.part.related_parts.values()  # type: ignore[union-attr]
    except Exception:
        logger.debug("doc.part.related_parts unavailable; skipping ZIP archive image count")
        return 0

    return sum(
        1
        for embedded_file in all_embedded_files
        if str(getattr(embedded_file, "content_type", "")).startswith("image/")
    )


# ---------------------------------------------------------------------------
# File opening
# ---------------------------------------------------------------------------


def open_word_document(path: Path) -> object:
    """
    Import python-docx and open the Word document at *path*.

    The import is done here rather than at the top of the file so that this
    module can load in environments where python-docx is not installed — the
    error only surfaces when extraction is actually attempted.

    Raises
    ------
    DocumentError
        If python-docx is not installed or the file cannot be opened.
    """
    try:
        from docx import Document  # noqa: PLC0415 — intentional deferred import
    except ImportError as exc:
        raise DocumentError(
            "python-docx is required for DOCX feature extraction. "
            "Install it with: pip install python-docx",
            context={"path": str(path)},
        ) from exc

    try:
        return Document(str(path))
    except Exception as exc:
        raise DocumentError(
            f"Could not open Word document for feature extraction: {path.name}",
            context={"path": str(path)},
        ) from exc
